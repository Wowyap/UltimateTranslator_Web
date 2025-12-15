import streamlit as st
import os
import io
import zipfile
from deep_translator import GoogleTranslator
from docx import Document
from pdf2docx import Converter
import tempfile
from collections import defaultdict
import uuid

# --- רשימת שפות ---
LANGUAGES = {
    'עברית': 'iw', 'אנגלית': 'en', 'רוסית': 'ru', 'ספרדית': 'es',
    'צרפתית': 'fr', 'ערבית': 'ar', 'גרמנית': 'de', 'איטלקית': 'it',
    'פורטוגזית': 'pt', 'סינית': 'zh-CN', 'יפנית': 'ja', 'הולנדית': 'nl'
}

# --- הגדרות דף ---
st.set_page_config(layout="wide", page_title="Ultimate Translator Web")
st.title("🌐 Ultimate Translator V5.0 - גרסת ענן")
st.markdown("מעבד קבצים במקביל (SRT, DOCX) וממיר/מתרגם PDF ל-DOCX")

# --- פונקציות מנוע התרגום (מותאמות לזיכרון) ---

@st.cache_resource
def get_translator(src_code, target_code):
    """יוצר אובייקט מתרגם ומטמין אותו"""
    return GoogleTranslator(source=src_code, target=target_code)

def trans_txt(file_bytes, tr):
    """תרגום קבצי טקסט (.vtt, .srt)"""
    lines = file_bytes.read().decode('utf-8').splitlines()
    new_lines = []
    
    for line in lines:
        t = line.strip()
        # מזהה שורות שאינן זמן/אינדקס/כותרת
        if "-->" not in t and t and not t.isdigit() and "WEBVTT" not in t:
            try:
                # מנסה לתרגם
                new_lines.append(tr.translate(t) + "\n")
            except:
                new_lines.append(line + "\n") # אם נכשל, שומר את המקור
        else:
            new_lines.append(line + "\n")
            
    # שומר ל-BytesIO כדי לשלוח כפלט
    output_buffer = io.BytesIO()
    output_buffer.write("".join(new_lines).encode('utf-8'))
    output_buffer.seek(0)
    return output_buffer

def trans_docx(file_bytes, tr):
    """תרגום קבצי Word (כולל טבלאות)"""
    doc = Document(file_bytes)
    
    # תרגום פסקאות
    for p in doc.paragraphs:
        if p.text.strip():
            try:
                p.text = tr.translate(p.text)
            except:
                pass
                
    # תרגום טבלאות
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text.strip():
                    try:
                        c.text = tr.translate(c.text)
                    except:
                        pass
                        
    # שמירה ל-BytesIO
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer

def trans_pdf(pdf_bytes, tr):
    """המרת PDF ל-DOCX, תרגום ושמירה"""
    
    # עבודה בתיקייה זמנית בשרת (חובה ל-pdf2docx)
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_pdf_path = os.path.join(temp_dir, f"input_{uuid.uuid4().hex[:8]}.pdf")
        temp_docx_path = os.path.join(temp_dir, f"output_{uuid.uuid4().hex[:8]}.docx")
        
        # 1. שמירת הקובץ שהועלה לנתיב זמני
        with open(temp_pdf_path, 'wb') as f:
            f.write(pdf_bytes.read())
            
        # 2. המרה מ-PDF ל-DOCX
        try:
            cv = Converter(temp_pdf_path)
            cv.convert(temp_docx_path, start=0, end=None)
            cv.close()
        except Exception as e:
            raise Exception(f"שגיאת המרה PDF ל-DOCX: {e}")
            
        # 3. קריאת ה-DOCX הזמני לזיכרון
        with open(temp_docx_path, 'rb') as f:
            docx_bytes = io.BytesIO(f.read())
            
        # 4. תרגום ה-DOCX
        translated_docx_buffer = trans_docx(docx_bytes, tr)
        
        return translated_docx_buffer


def process_file_in_memory(uploaded_file, tr):
    """הפונקציה המרכזית לעיבוד קובץ יחיד"""
    
    filename = uploaded_file.name
    ext = os.path.splitext(filename)[1].lower()
    file_bytes = uploaded_file
    
    if ext in ['.vtt', '.srt']:
        translated_buffer = trans_txt(file_bytes, tr)
        new_ext = ext
        
    elif ext == '.docx':
        translated_buffer = trans_docx(file_bytes, tr)
        new_ext = '.docx'
        
    elif ext == '.pdf':
        translated_buffer = trans_pdf(file_bytes, tr)
        new_ext = '.docx' # פלט PDF הוא תמיד DOCX
        
    else:
        raise Exception(f"סוג קובץ לא נתמך: {ext}")
        
    # יצירת שם קובץ חדש (שם מקורי + קוד שפה)
    base_name = os.path.splitext(filename)[0]
    new_filename = f"{base_name}.{tr.target_language}{new_ext}"
    
    return new_filename, translated_buffer

# --- ממשק משתמש וניהול תהליך ---

with st.sidebar:
    st.header("1. הגדרות שפה")
    col_src, col_target = st.columns(2)
    
    with col_src:
        src_lang = st.selectbox("שפת מקור:", ['זיהוי אוטומטי'] + list(LANGUAGES.keys()), index=0)
    
    with col_target:
        target_lang = st.selectbox("שפת יעד:", list(LANGUAGES.keys()), index=LANGUAGES.keys().index('עברית'))
        
    st.header("2. קבצים")
    uploaded_files = st.file_uploader(
        "גרור והעלה קבצים (DOCX, PDF, SRT, VTT)", 
        type=['docx', 'pdf', 'srt', 'vtt'], 
        accept_multiple_files=True
    )
    
    if len(uploaded_files) > 0:
        st.info(f"סה\"כ {len(uploaded_files)} קבצים מוכנים.")

if uploaded_files:
    
    if st.button("🚀 התחל תרגום קבצים"):
        
        if src_lang == 'זיהוי אוטומטי':
            src_code = 'auto'
        else:
            src_code = LANGUAGES[src_lang]
            
        target_code = LANGUAGES[target_lang]
        
        # אתחול מתרגם
        tr = get_translator(src_code, target_code)
        
        st.subheader("🚧 סטטוס עיבוד")
        progress_bar = st.progress(0, text="מתחיל...")
        
        # משתנים לאחסון התוצאות
        translated_files = []
        errors = defaultdict(list)
        
        for i, file in enumerate(uploaded_files):
            file_name = file.name
            
            try:
                # Streamlit קורא את הקובץ לזיכרון וסוגר אותו, צריך לקרוא אותו שוב
                file.seek(0)
                
                # העיבוד בפועל
                new_filename, buffer = process_file_in_memory(file, tr)
                translated_files.append((new_filename, buffer))
                st.success(f"✅ הושלם: {file_name} -> {new_filename}")
                
            except Exception as e:
                errors['שגיאה'].append(f"{file_name}: {e}")
                st.error(f"❌ שגיאה בקובץ {file_name}")

            # עדכון סרגל התקדמות
            progress_bar.progress((i + 1) / len(uploaded_files), text=f"מעבד: {file_name} ({i+1}/{len(uploaded_files)})")

        
        progress_bar.empty()
        
        if translated_files:
            st.subheader("📥 תוצאות להורדה")
            
            # יצירת קובץ ZIP להורדה
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                for new_name, buffer in translated_files:
                    # שימוש ב-buffer.getvalue() כדי לקבל את ה-bytes
                    zip_file.writestr(new_name, buffer.getvalue())
            
            zip_buffer.seek(0)

            st.download_button(
                label=f"הורד את כל {len(translated_files)} הקבצים (ZIP)",
                data=zip_buffer,
                file_name=f"Translated_Files_{target_code}.zip",
                mime="application/zip"
            )

        if errors:
            st.subheader("🛑 סיכום שגיאות")
            st.error("\n".join(errors['שגיאה']))